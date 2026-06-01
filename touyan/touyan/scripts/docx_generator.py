#!/usr/bin/env python3
"""
touyan 研判报告 docx 自动生成脚本
依赖：pip install python-docx
用法：python docx_generator.py --company "公司名" --date "2026-05-27" --output /path/to/output.docx --data /path/to/data.json
"""

import json
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import sys
import os

# ============ 字体工具 ============

def set_font(run, cn_font, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def set_line_spacing(paragraph, spacing_pt=28):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = pPr.makeelement(qn('w:spacing'), {})
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(spacing_pt * 20))
    spacing.set(qn('w:lineRule'), 'exact')

def set_first_indent(paragraph, chars=2, font_size_pt=16.5):
    paragraph.paragraph_format.first_line_indent = Pt(font_size_pt * chars)

def set_no_indent(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)

# ============ 文档构建工具 ============

def new_para(doc, text, cn_font='仿宋_GB2312', size=16.5, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, line_spacing=28):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, cn_font, size, bold)
    set_line_spacing(p, line_spacing)
    p.alignment = align
    if indent:
        set_first_indent(p, 2, size)
    else:
        set_no_indent(p)
    return p

def add_title_lines(doc, company_name):
    """公文标题区域"""
    # 第一行
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("成都诸葛私募基金管理有限公司关于")
    set_font(run1, '方正公文小标宋', 22)
    set_line_spacing(p1, 28)
    # 第二行
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{company_name}\n项目的初步研判")
    set_font(run2, '方正公文小标宋', 22)
    set_line_spacing(p2, 28)

def add_core_verdict(doc, verdict_text):
    """核心判断区（加粗，一到两句话）"""
    p = doc.add_paragraph()
    run = p.add_run(verdict_text)
    set_font(run, '仿宋_GB2312', 16.5, bold=True)
    set_line_spacing(p, 28)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_recipient_line(doc, text="晋阳街道："):
    """主送行，顶格"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5)
    set_line_spacing(p, 28)
    set_no_indent(p)
    return p

def add_opening(doc, text):
    """开头段"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5)
    set_line_spacing(p, 28)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_h1(doc, text):
    """一级标题（一、）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '黑体', 16.5)
    set_line_spacing(p, 28)
    set_first_indent(p, 2, 16.5)
    return p

def add_h2(doc, text):
    """二级标题（（一））"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '楷体_GB2312', 16.5, bold=True)
    set_line_spacing(p, 28)
    set_first_indent(p, 2, 16.5)
    return p

def add_h3(doc, text):
    """三级标题（1.）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5, bold=True)
    set_line_spacing(p, 28)
    set_first_indent(p, 2, 16.5)
    return p

def add_body(doc, text):
    """正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5)
    set_line_spacing(p, 28)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_first_indent(p, 2, 16.5)
    return p

def add_table(doc, headers, rows_data):
    """表格：表头黑体居中，数据行仿宋居中"""
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, '黑体', 16.5)
        set_line_spacing(p, 28)
    # 数据行
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, '仿宋_GB2312', 16.5)
            set_line_spacing(p, 28)
    return table

def add_signature(doc, company="成都诸葛私募基金管理有限公司", date_text="YYYY年MM月DD日"):
    """落款区"""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = p1.add_run(company)
    set_font(run1, '仿宋_GB2312', 16.5)
    set_line_spacing(p1, 28)
    set_no_indent(p1)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(date_text)
    set_font(run2, '仿宋_GB2312', 16.5)
    set_line_spacing(p2, 28)
    set_no_indent(p2)

def brush_western_font(doc):
    """最后一步：遍历全文确保西文字体为Times New Roman"""
    for para in doc.paragraphs:
        for run in para.runs:
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        r = run._element
                        rPr = r.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                rFonts.set(qn('w:ascii'), 'Times New Roman')
                                rFonts.set(qn('w:hAnsi'), 'Times New Roman')

# ============ 主生成函数 ============

def build_report(data):
    """
    data 格式（JSON）：
    {
        "company": "公司全称",
        "date": "YYYY年MM月DD日",
        "core_verdict": "核心判断文字",
        "opening": "开头段文字",
        "sections": [
            {
                "h1": "一、基本情况",
                "h2s": [
                    {"h2": "（一）工商信息", "body": "正文内容..."},
                    {"h2": "（二）股权结构...", "body": "..."}
                ]
            },
            {
                "h1": "二、初步研判",
                "h2s": [
                    {"h2": "（一）行业研判", "body": "..."},
                    {"h2": "（二）公司研判", "body": "..."},
                    {"h2": "（三）估值与交易结构", "body": "..."}
                ]
            },
            {
                "h1": "三、综合研判",
                "body": "综合研判段落..."
            }
        ],
        "confidence_table": {
            "enabled": true,
            "headers": ["关键数据点", "来源", "置信度", "需核实"],
            "rows": [["XXX", "企业自述", "低", "待核实"]]
        },
        "signature": "成都诸葛私募基金管理有限公司"
    }
    """
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    company = data.get("company", "标的企业")
    date_text = data.get("date", "YYYY年MM月DD日")
    core_verdict = data.get("core_verdict", "")
    opening = data.get("opening", "")
    sections = data.get("sections", [])
    has_confidence_table = data.get("confidence_table", {}).get("enabled", False)
    confidence_rows = data.get("confidence_table", {}).get("rows", [])
    sig = data.get("signature", "成都诸葛私募基金管理有限公司")

    # 标题区
    add_title_lines(doc, company)

    # 核心判断（加粗）
    if core_verdict:
        add_core_verdict(doc, core_verdict)

    # 主送
    add_recipient_line(doc)

    # 开头段
    if opening:
        add_opening(doc, opening)

    # 主体章节
    for sec in sections:
        if "h1" in sec:
            add_h1(doc, sec["h1"])
        if "h2s" in sec:
            for h2_item in sec["h2s"]:
                if "h2" in h2_item:
                    add_h2(doc, h2_item["h2"])
                if "body" in h2_item and h2_item["body"]:
                    add_body(doc, h2_item["body"])
        if "body" in sec and sec["body"]:
            add_body(doc, sec["body"])

    # 置信度评分卡（仅当启用时）
    if has_confidence_table and confidence_rows:
        add_h2(doc, "（附）信息置信度评分卡")
        add_table(doc,
                 data.get("confidence_table", {}).get("headers", []),
                 confidence_rows)

    # 落款
    add_signature(doc, sig, date_text)

    # 最后刷一遍西文字体
    brush_western_font(doc)

    return doc


# ============ CLI 入口 ============

def main():
    parser = argparse.ArgumentParser(description='touyan 研判报告 docx 生成器')
    parser.add_argument('--company', '-c', required=True, help='企业全称')
    parser.add_argument('--date', '-d', default='YYYY年MM月DD日', help='日期')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--data', '-j', help='JSON数据文件路径（替代命令行逐项输入）')
    parser.add_argument('--verdict', '-v', default='', help='核心判断')
    parser.add_argument('--opening', help='开头段文字（文件路径或直接文本）')

    args = parser.parse_args()

    # 如果有JSON文件，优先使用
    if args.data and os.path.exists(args.data):
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        # 从参数构建最小结构（用户需在 JSON 中提供完整 sections）
        print("ERROR: 目前需要通过 --data 指定 JSON 文件。")
        print("请使用以下格式准备 JSON 文件：")
        sample = {
            "company": args.company,
            "date": args.date,
            "core_verdict": args.verdict,
            "opening": "根据...，现形成初步研判报告如下。",
            "sections": [
                {"h1": "一、基本情况", "h2s": [
                    {"h2": "（一）工商信息", "body": "..."}
                ]}
            ],
            "confidence_table": {"enabled": False, "headers": [], "rows": []}
        }
        print(json.dumps(sample, ensure_ascii=False, indent=2))
        sys.exit(1)

    doc = build_report(data)
    doc.save(args.output)
    print(f"✅ 报告已生成：{args.output}")


if __name__ == '__main__':
    main()
