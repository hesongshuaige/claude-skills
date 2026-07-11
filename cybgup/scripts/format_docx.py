#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GB/T 9704-2012 公文排版：纯文本 → Word(.docx)
逐行识别结构（标题/抬头/各级标题/正文/落款），按国标排版，不改内容只管格式。

用法:
  python format_docx.py input.txt -o output.docx        # 纯文本
  python format_docx.py input.md -o output.docx --md    # 先剥离 Markdown 标记
  python format_docx.py - -o output.docx                # 从 stdin 读
"""
import argparse
import re
import sys

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# ============ 字体与字号（GB/T 9704-2012）============
F_TITLE = "方正公文小标宋"   # 公文标题（二号 22pt，居中，不加粗）
F_L1 = "黑体"                # 一级标题 一、（16.5pt，不加粗）
F_L2 = "楷体_GB2312"         # 二级标题 （一）（16.5pt，加粗）
F_FANGSONG = "仿宋_GB2312"   # 三级标题/正文/抬头/保密/落款（16.5pt）
F_WEST = "Times New Roman"   # 西文统一
SZ_TITLE = 22                # 二号
SZ_BODY = 16.5
LINE_PT = 28                 # 行距固定值 28 磅
INDENT_PT = 33               # 首行缩进 2 字符 = 2 × 16.5pt

# ============ 结构识别正则 ============
RE_L1 = re.compile(r"^[一二三四五六七八九十]+\s*、")
RE_L2 = re.compile(r"^[（(]\s*[一二三四五六七八九十]+\s*[）)]")
RE_L3 = re.compile(r"^\d+\s*[.、]")
RE_SALUT = re.compile(r"^.{1,50}[：:]$")
RE_DATE = re.compile(r"20\d{2}年\d{1,2}月\d{1,2}日\s*$")
RE_SECRET = re.compile(r"^【.*】$")
RE_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|\s*$")
TITLE_BAD = set("，。；！？：")
SENT_END = set("。！？；")


def strip_markdown(text: str) -> str:
    """剥离 Markdown 标记转纯文本（frontmatter / # 标题 / **粗体** / 链接 / 表格管道）。"""
    out, in_fm, fm_seen = [], False, 0
    for ln in text.splitlines():
        s = ln.rstrip()
        stripped = s.strip()
        if stripped == "---":                       # frontmatter 围栏
            fm_seen += 1
            in_fm = fm_seen == 1
            continue
        if in_fm:
            continue
        if RE_TABLE_SEP.match(stripped):            # |---|---| 分隔行
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)  # ATX 标题前缀
        if m:
            stripped = m.group(2)
        stripped = stripped.replace("**", "").replace("__", "")  # 粗体
        stripped = re.sub(r"(?<!\*)\*(?!\*)", "", stripped)       # 斜体 *
        stripped = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", stripped)  # 链接留文本
        if stripped.startswith("|") and stripped.endswith("|"):    # 表格行：拼成一行
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            stripped = "　".join(cells)
        out.append(stripped)
    return "\n".join(out)


def is_date_line(s: str) -> bool:
    return bool(RE_DATE.search(s)) and len(s) <= 30


def classify(lines):
    """逐行分类，返回 [(type, text), ...]，type 见 STYLE。"""
    nonblank = [ln.strip() for ln in lines if ln.strip() != ""]
    res, title_phase, i, n = [], True, 0, len(nonblank)
    while i < n:
        s = nonblank[i]
        # 1. 保密标记
        if RE_SECRET.match(s):
            title_phase = False
            res.append(("secret", s)); i += 1; continue
        # 2-4. 各级标题
        if RE_L1.match(s):
            title_phase = False; res.append(("l1", s)); i += 1; continue
        if RE_L2.match(s):
            title_phase = False; res.append(("l2", s)); i += 1; continue
        if RE_L3.match(s):
            title_phase = False; res.append(("l3", s)); i += 1; continue
        # 5. 抬头/主送
        if RE_SALUT.match(s):
            title_phase = False; res.append(("salutation", s)); i += 1; continue
        # 6. 落款日期
        if is_date_line(s):
            title_phase = False; res.append(("signature", s)); i += 1; continue
        # 8. 落款单位名（前瞻：下一行是日期）
        if i + 1 < n and is_date_line(nonblank[i + 1]) and len(s) <= 30 \
                and not s[-1] in SENT_END:
            title_phase = False; res.append(("signature", s)); i += 1; continue
        # 7. 公文标题（开头连续短行、无标点）
        if title_phase and len(s) <= 50 and not any(c in TITLE_BAD for c in s):
            res.append(("title", s)); i += 1; continue
        # 9. 正文
        title_phase = False
        res.append(("body", s)); i += 1
    return res


def insert_blanks(items):
    """按规则插空行：标题后空一行；落款前空两行（仅首次）。"""
    last_title = -1
    for idx, (t, _) in enumerate(items):
        if t == "title":
            last_title = idx
        elif last_title >= 0:
            break
    first_sig = next((i for i, (t, _) in enumerate(items) if t == "signature"), None)
    out = []
    for idx, (t, txt) in enumerate(items):
        if idx == first_sig:
            out.append(("blank", "")); out.append(("blank", ""))
        out.append((t, txt))
        if idx == last_title:
            out.append(("blank", ""))
    return out


# 每类的样式：(中文字体, 字号pt, 加粗, 对齐, 首行缩进)
STYLE = {
    "title":      (F_TITLE,    SZ_TITLE, False, WD_ALIGN_PARAGRAPH.CENTER,  False),
    "l1":         (F_L1,       SZ_BODY,  False, WD_ALIGN_PARAGRAPH.JUSTIFY, True),
    "l2":         (F_L2,       SZ_BODY,  True,  WD_ALIGN_PARAGRAPH.JUSTIFY, True),
    "l3":         (F_FANGSONG, SZ_BODY,  False, WD_ALIGN_PARAGRAPH.JUSTIFY, True),
    "body":       (F_FANGSONG, SZ_BODY,  False, WD_ALIGN_PARAGRAPH.JUSTIFY, True),
    "salutation": (F_FANGSONG, SZ_BODY,  False, WD_ALIGN_PARAGRAPH.LEFT,    False),
    "secret":     (F_FANGSONG, SZ_BODY,  True,  WD_ALIGN_PARAGRAPH.CENTER,  False),
    "signature":  (F_FANGSONG, SZ_BODY,  False, WD_ALIGN_PARAGRAPH.RIGHT,   False),
    "blank":      (F_FANGSONG, SZ_BODY,  False, WD_ALIGN_PARAGRAPH.LEFT,    False),
}


def style_run(run, cn_font, size_pt, bold):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = F_WEST
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), F_WEST)
    rfonts.set(qn("w:hAnsi"), F_WEST)
    rfonts.set(qn("w:eastAsia"), cn_font)


def fmt_para(p, kind, text):
    cn_font, size_pt, bold, align, indent = STYLE[kind]
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(LINE_PT)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Pt(INDENT_PT)
    run = p.add_run(text)
    style_run(run, cn_font, size_pt, bold)


def build_docx(items, out_path):
    doc = Document()
    # 页面：A4 + 国标页边距
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    # Normal 默认字体，兜底空段落
    normal = doc.styles["Normal"]
    normal.font.name = F_WEST
    normal.font.size = Pt(SZ_BODY)
    for kind, txt in items:
        fmt_para(doc.add_paragraph(), kind, txt)
    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser(description="GB/T 9704-2012 公文排版：纯文本 → Word")
    ap.add_argument("input", help="输入文件路径，或 - 读 stdin")
    ap.add_argument("-o", "--output", required=True, help="输出 .docx 路径")
    ap.add_argument("--md", action="store_true", help="先剥离 Markdown 标记")
    args = ap.parse_args()

    if args.input == "-":
        raw = sys.stdin.read().lstrip("\ufeff")
    else:
        with open(args.input, encoding="utf-8-sig") as f:
            raw = f.read()
    if args.md:
        raw = strip_markdown(raw)
    items = insert_blanks(classify(raw.splitlines()))
    build_docx(items, args.output)
    # 简报：各类计数
    from collections import Counter
    cnt = Counter(t for t, _ in items)
    print(f"OK -> {args.output}")
    print("结构统计: " + ", ".join(f"{k}={v}" for k, v in cnt.items()))


if __name__ == "__main__":
    main()