#!/usr/bin/env python3
"""投资研判报告质量门禁脚本

检查报告草稿是否包含所有硬性要素。
门禁只检查要素是否出现，不替代事实核查和人工判断。
门禁失败时必须补齐报告或在信息缺失表中说明原因。

用法：
  python3 quality_gate.py report.md
  python3 quality_gate.py report.md --formal --national-assets
  python3 quality_gate.py report.md --json  # JSON格式输出，方便脚本集成
"""

import argparse
import json
import re
import sys
from pathlib import Path


# ============ 工具函数 ============

def flatten_json(value):
    """递归展平JSON结构为纯文本"""
    if isinstance(value, dict):
        return "\n".join(flatten_json(v) for v in value.values())
    if isinstance(value, list):
        return "\n".join(flatten_json(v) for v in value)
    return "" if value is None else str(value)


def read_text(path):
    """读取文件内容，支持 .md/.txt/.json/.docx"""
    suffix = path.suffix.lower()
    if suffix == ".json":
        return flatten_json(json.loads(path.read_text(encoding="utf-8")))
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to inspect docx files") from exc
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8")


def has_any(text, terms):
    """检查文本中是否包含任一关键词"""
    return any(term in text for term in terms)


def has_verdict(text):
    """检查报告是否包含明确结论（精确匹配，防止"推进中"等过程描述误通过）

    只接受三种结论性用语：
    - "不投" / "建议不投"
    - "有条件推进" / "建议有条件推进"
    - "建议推进" / "可以推进"

    排除："推进中" "正在推进" "持续推进" 等过程描述
    """
    # 匹配结论性表述
    # 关键：结论词后面不能紧跟"中"/"过程"等过程描述词
    # 使用负向前瞻确保"推进"后面不是"中"/"过程"/"工作"等
    verdict_patterns = [
        r'(?:建议)?不投(?!资)',                              # "不投"但不能是"不投资"
        r'(?:建议)?有条件推进(?!中|过程|工作|阶段)',          # "有条件推进"后不能跟过程词
        r'建议推进(?!中|过程|工作|阶段)',                     # "建议推进"后不能跟过程词
        r'可以推进(?!中|过程|工作|阶段)',                     # "可以推进"后不能跟过程词
        r'结论[：:]\s*(?:不投|有条件推进|建议推进)',
        r'综合研判[，,]?\s*(?:不投|有条件推进|建议推进)',
    ]
    for pattern in verdict_patterns:
        if re.search(pattern, text):
            return True
    return False


def has_weighted_dpi(text):
    """检查是否包含概率加权DPI（兼容多种表述）"""
    return has_any(text, ["概率加权", "加权综合", "加权平均DPI", "概率加权DPI"])


# ============ 检查项定义 ============

# 每个检查项：(名称, 检查函数, 失败提示)
def get_base_checks():
    """基础检查项（所有报告必须通过）"""
    return [
        ("明确结论（不投/有条件推进/建议推进）",
         has_verdict(text_global),
         "报告必须包含'不投''有条件推进'或'建议推进'作为明确结论，不接受'推进中'等过程描述"),

        ("信息缺失分级 P0/P1/P2",
         all(term in text_global for term in ["P0", "P1", "P2"]),
         "报告必须包含P0/P1/P2信息缺失分级表，即使无缺失也须注明'无'"),

        ("置信度标注",
         "置信度" in text_global,
         "报告中的关键数据必须标注置信度等级（高/中/低/缺失）"),

        ("五情景DPI表（乐观/基准/保守/回购/清算）",
         all(term in text_global for term in ["乐观", "基准", "保守", "回购", "清算", "DPI"]),
         "退出分析必须包含五情景DPI表：乐观/基准/保守/回购/清算+DPI"),

        ("概率加权综合DPI",
         has_weighted_dpi(text_global),
         "DPI表必须包含概率加权综合DPI"),

        ("交易条款推导链路",
         has_any(text_global, ["推导", "依据", "测算逻辑", "参照"]),
         "交易条款的数字必须有推导依据（如'参照XX数据/按XX逻辑推算'）"),

        ("报告日期",
         has_any(text_global, ["年", "月", "日"]) and len(re.findall(r'\d{4}', text_global)) > 0,
         "报告必须包含日期"),
    ]


def get_national_assets_checks():
    """国资合规检查项（--national-assets 时启用）"""
    checks = [
        ("12号令（资产评估）",
         "12号令" in text_global,
         "涉及国资出资时须引用《企业国有资产评估管理暂行办法》（12号令）"),

        ("32号令（进场交易）",
         "32号令" in text_global,
         "涉及国资交易时须引用《企业国有资产交易监督管理办法》（32号令）"),

        ("明股实债风险",
         "明股实债" in text_global,
         "须核查是否存在明股实债风险"),

        ("变相保底风险",
         "变相保底" in text_global,
         "须核查是否存在变相保底承诺"),

        ("关联交易",
         "关联交易" in text_global,
         "须核查关联交易情况"),
    ]

    # 涉及上市公司国有股权时，需引用36号令
    if has_any(text_global, ["上市公司", "国有股权"]):
        checks.append((
            "36号令（上市公司国有股权）",
            "36号令" in text_global,
            "涉及上市公司国有股权时须引用36号令，不得与32号令混用"
        ))

    return checks


def get_formal_checks():
    """正式尽调检查项（--formal 时启用）"""
    return [
        ("正式尽调-财务",
         "财务" in text_global,
         "正式尽调须覆盖财务专项分析"),

        ("正式尽调-商业",
         has_any(text_global, ["商业", "客户验证", "竞争格局"]),
         "正式尽调须覆盖商业/客户/竞争分析"),

        ("正式尽调-法律治理",
         has_any(text_global, ["法律", "治理"]),
         "正式尽调须覆盖法律与治理分析"),

        ("正式尽调-技术IP",
         has_any(text_global, ["技术", "IP", "知识产权", "专利"]),
         "正式尽调须覆盖技术/IP分析"),

        ("基金回报指标（MOIC/IRR/DPI）",
         has_any(text_global, ["MOIC", "IRR", "DPI"]),
         "正式尽调须包含MOIC/IRR/DPI等基金回报指标"),

        ("不能替代正式尽调声明",
         has_any(text_global, ["不能替代正式尽调", "不能替代正式投决尽调"]),
         "若核心证据缺失，须明确声明'本报告不能替代正式尽调'"),
    ]


# ============ 主流程 ============

# 全局变量，用于在检查函数中访问文本
text_global = ""


def run_checks(text, formal=False, national_assets=False):
    """运行所有检查项，返回 (checks_with_results, failed_list)

    Args:
        text: 报告文本内容
        formal: 是否启用正式尽调检查
        national_assets: 是否启用国资合规检查

    Returns:
        (all_checks, failed_names)
        all_checks: [(name, passed, hint), ...]
        failed_names: [name, ...] 未通过的检查项名称
    """
    global text_global
    text_global = text

    checks = get_base_checks()

    if national_assets:
        checks.extend(get_national_assets_checks())

    if formal:
        checks.extend(get_formal_checks())

    failed = []
    for item in checks:
        name = item[0]
        passed = item[1]
        hint = item[2] if len(item) > 2 else ""
        if not passed:
            failed.append((name, hint))

    return checks, failed


def main():
    parser = argparse.ArgumentParser(description="投资研判报告质量门禁 — 检查硬性要素是否齐全")
    parser.add_argument("report", help="报告文件路径：.md, .txt, .json, 或 .docx")
    parser.add_argument("--formal", action="store_true", help="启用正式PE尽调检查项")
    parser.add_argument("--national-assets", action="store_true", help="启用国资合规检查项")
    parser.add_argument("--json", action="store_true", help="JSON格式输出（方便脚本集成）")
    args = parser.parse_args()

    path = Path(args.report)
    if not path.exists():
        msg = f"文件不存在: {path}"
        if args.json:
            print(json.dumps({"passed": False, "error": msg}, ensure_ascii=False))
        else:
            print(f"ERROR: {msg}", file=sys.stderr)
        return 2

    text = read_text(path)
    checks, failed = run_checks(text, formal=args.formal, national_assets=args.national_assets)

    # JSON输出
    if args.json:
        results = []
        for item in checks:
            results.append({
                "name": item[0],
                "passed": bool(item[1]),
                "hint": item[2] if len(item) > 2 else ""
            })
        output = {
            "passed": len(failed) == 0,
            "total": len(checks),
            "failed_count": len(failed),
            "results": results,
            "failed": [{"name": f[0], "hint": f[1]} for f in failed]
        }
        print(json.dumps(output, ensure_ascii=False, indent=2))
        return 0 if not failed else 1

    # 文本输出
    for item in checks:
        status = "PASS" if item[1] else "FAIL"
        print(f"  {status}: {item[0]}")

    if failed:
        print(f"\n❌ 未通过 {len(failed)} 项检查：")
        for name, hint in failed:
            print(f"  ✗ {name}")
            if hint:
                print(f"    → {hint}")
        return 1

    print("\n✅ 所有硬性检查通过。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
