#!/usr/bin/env python3
"""
投资研判报告 docx 自动生成脚本

将JSON格式的报告数据生成符合公文排版标准的Word文档。
依赖：pip install python-docx

用法：
  python3 docx_generator.py --company "公司名" --date "2026-06-02" --data report.json --output output.docx
"""

import json
import argparse
import sys
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============ 数据验证 ============

def validate_data(data):
    """验证JSON数据结构完整性

    Args:
        data: 解析后的JSON数据(dict)

    Returns:
        list[str]: 错误信息列表，空列表表示验证通过
    """
    errors = []

    if not data.get("company", "").strip():
        errors.append("缺少必填字段: company（企业全称）")

    sections = data.get("sections", [])
    if not sections:
        errors.append("缺少必填字段: sections（报告主体章节，不能为空列表）")
    else:
        for i, sec in enumerate(sections):
            if not sec.get("h1") and not sec.get("body"):
                errors.append(f"sections[{i}] 缺少 h1 或 body，至少需要其一")

    return errors


# ============ 字体工具 ============

def set_font(run, cn_font, size_pt, bold=False):
    """设置run的中英文字体、字号、加粗

    Args:
        run: python-docx Run对象
        cn_font: 中文字体名称
        size_pt: 字号（磅）
        bold: 是否加粗
    """
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def set_line_spacing(paragraph, spacing_pt=28):
    """设置段落固定行距

    Args:
        paragraph: 段落对象
        spacing_pt: 行距值（磅），公文标准28pt
    """
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = pPr.makeelement(qn('w:spacing'), {})
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(spacing_pt * 20))
    spacing.set(qn('w:lineRule'), 'exact')


def set_first_indent(paragraph, chars=2, font_size_pt=16.5):
    """设置首行缩进

    Args:
        paragraph: 段落对象
        chars: 缩进字符数
        font_size_pt: 字号（用于计算缩进量）
    """
    paragraph.paragraph_format.first_line_indent = Pt(font_size_pt * chars)


def set_no_indent(paragraph):
    """取消缩进（用于标题、落款等）"""
    paragraph.paragraph_format.first_line_indent = Pt(0)


# ============ 文档构建工具 ============

def new_para(doc, text, cn_font='仿宋_GB2312', size=16.5, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, line_spacing=28):
    """通用段落创建函数"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, cn_font, size, bold)
    set_line_spacing(p, line_spacing)
    p.alignment = align
    if indent:
        set_first_indent(p, 2, size)
    else:
        set_no_indent(p)
    return p


def add_title_lines(doc, company_name, fund_name="投资机构"):
    """公文标题区域（两行居中标题）"""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(f"{fund_name}关于")
    set_font(run1, '方正公文小标宋', 22)
    set_line_spacing(p1, 28)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{company_name}\n项目的初步研判")
    set_font(run2, '方正公文小标宋', 22)
    set_line_spacing(p2, 28)


def add_core_verdict(doc, verdict_text):
    """核心判断区（加粗，让领导一眼看到结论）"""
    p = doc.add_paragraph()
    run = p.add_run(verdict_text)
    set_font(run, '仿宋_GB2312', 16.5, bold=True)
    set_line_spacing(p, 28)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_recipient_line(doc, text="主送单位："):
    """主送行（顶格）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5)
    set_line_spacing(p, 28)
    set_no_indent(p)
    return p


def add_opening(doc, text):
    """开头段"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5)
    set_line_spacing(p, 28)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_h1(doc, text):
    """一级标题（一、）— 黑体"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '黑体', 16.5)
    set_line_spacing(p, 28)
    set_first_indent(p, 2, 16.5)
    return p


def add_h2(doc, text):
    """二级标题（（一））— 楷体加粗"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '楷体_GB2312', 16.5, bold=True)
    set_line_spacing(p, 28)
    set_first_indent(p, 2, 16.5)
    return p


def add_body(doc, text):
    """正文段落 — 仿宋"""
    if not text or not text.strip():
        return None
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '仿宋_GB2312', 16.5)
    set_line_spacing(p, 28)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_first_indent(p, 2, 16.5)
    return p


def set_cell_shading(cell, color="D9E2F3"):
    """设置单元格背景色（用于表头高亮）

    Args:
        cell: 表格单元格
        color: 十六进制颜色值，默认浅蓝
    """
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def set_table_header_repeat(table):
    """设置表格第一行（表头）跨页时自动重复

    Args:
        table: python-docx Table对象
    """
    if not table.rows:
        return
    first_row = table.rows[0]
    for cell in first_row.cells:
        tr = cell._element.getparent()
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = tr.makeelement(qn('w:trPr'), {})
            tr.insert(0, trPr)
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = trPr.makeelement(qn('w:tblHeader'), {})
            trPr.append(tblHeader)


def add_table(doc, headers, rows_data):
    """添加表格：表头黑体居中+浅蓝背景，数据行仿宋居中

    Args:
        doc: Document对象
        headers: 表头列表
        rows_data: 数据行列表（二维列表）

    Returns:
        Table对象
    """
    if not headers:
        return None

    row_count = max(len(rows_data), 0) + 1
    table = doc.add_table(rows=row_count, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, '黑体', 16.5)
        set_line_spacing(p, 28)
        set_cell_shading(cell, "D9E2F3")

    # 数据行
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, '仿宋_GB2312', 16.5)
            set_line_spacing(p, 28)

    # 设置表头跨页重复
    set_table_header_repeat(table)

    return table


def add_page_break(doc):
    """插入分页符（用于置信度表、短信附件等前分页）"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)


def add_signature(doc, company="投资机构", date_text="YYYY年MM月DD日"):
    """落款区（右对齐）"""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = p1.add_run(company)
    set_font(run1, '仿宋_GB2312', 16.5)
    set_line_spacing(p1, 28)
    set_no_indent(p1)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(date_text)
    set_font(run2, '仿宋_GB2312', 16.5)
    set_line_spacing(p2, 28)
    set_no_indent(p2)


def brush_western_font(doc):
    """最后一步：遍历全文确保西文字体统一为Times New Roman"""
    targets = []
    targets.extend(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)

    for para in targets:
        for run in para.runs:
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


# ============ 主生成函数 ============

def build_report(data):
    """从JSON数据构建完整的公文格式docx

    Args:
        data: 符合schema的JSON数据(dict)

    Returns:
        Document对象
    """
    doc = Document()

    # 页面设置（公文标准）
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    company = data.get("company", "标的企业")
    date_text = data.get("date", "YYYY年MM月DD日")
    core_verdict = data.get("core_verdict", "")
    opening = data.get("opening", "")
    sections = data.get("sections", [])
    sig = data.get("signature", data.get("fund_name", "投资机构"))
    fund_name = data.get("fund_name", sig)
    recipient = data.get("recipient", "主送单位：")

    # 标题区
    add_title_lines(doc, company, fund_name)

    # 核心判断（加粗）
    if core_verdict:
        add_core_verdict(doc, core_verdict)

    # 主送
    add_recipient_line(doc, recipient)

    # 开头段
    if opening:
        add_opening(doc, opening)

    # 主体章节
    for sec in sections:
        if sec.get("h1"):
            add_h1(doc, sec["h1"])

        # 二级标题和正文
        h2s = sec.get("h2s", [])
        if h2s:
            for h2_item in h2s:
                if h2_item.get("h2"):
                    add_h2(doc, h2_item["h2"])
                if h2_item.get("body") and h2_item["body"].strip():
                    add_body(doc, h2_item["body"])

        # 无二级标题时的正文
        if sec.get("body") and sec["body"].strip():
            add_body(doc, sec["body"])

    # 置信度评分卡（仅当启用且有数据时）
    conf = data.get("confidence_table", {})
    if conf.get("enabled") and conf.get("rows"):
        add_h2(doc, "（附）信息置信度评分卡")
        add_table(doc, conf.get("headers", []), conf["rows"])

    # 短信附件（位于落款之前）
    sms_appendix = data.get("sms_appendix", {})
    if sms_appendix.get("enabled") and sms_appendix.get("text"):
        sms_date = sms_appendix.get("date", date_text)
        add_h1(doc, f"附件：领导简报短信（{sms_date}）")
        # 术语差异说明
        add_body(doc, "（注：以下短信版本已将DPI/MOIC/IRR等专业术语转换为通俗表述，与报告正文含义一致，仅表述方式不同。供直接转发使用。）")
        # 按段落分割短信文本
        sms_text = sms_appendix["text"]
        for para_text in sms_text.split("\n"):
            para_text = para_text.strip()
            if para_text:
                add_body(doc, para_text)

    # 落款
    add_signature(doc, sig, date_text)

    # 最后刷一遍西文字体
    brush_western_font(doc)

    return doc


# ============ CLI 入口 ============

def main():
    parser = argparse.ArgumentParser(description='投资研判报告 docx 生成器（公文排版标准）')
    parser.add_argument('--company', '-c', required=True, help='企业全称')
    parser.add_argument('--date', '-d', default='YYYY年MM月DD日', help='日期')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--data', '-j', required=True, help='JSON数据文件路径')
    args = parser.parse_args()

    # 读取JSON数据
    if not os.path.exists(args.data):
        print(f"ERROR: JSON文件不存在: {args.data}", file=sys.stderr)
        sys.exit(1)

    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 数据验证
    errors = validate_data(data)
    if errors:
        print("❌ 数据验证失败：", file=sys.stderr)
        for e in errors:
            print(f"  • {e}", file=sys.stderr)
        sys.exit(1)

    # 输出目录自动创建
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # 生成文档
    doc = build_report(data)
    doc.save(args.output)
    print(f"✅ 报告已生成：{args.output}")


if __name__ == '__main__':
    main()
